"""
Document processing module with extensible architecture.
Handles multiple document formats (PDF, DOC, DOCX) with text extraction and chunking.
"""

import logging
from pathlib import Path
from typing import List, Tuple, Protocol
from pypdf import PdfReader
from docx import Document
import re

logger = logging.getLogger(__name__)


class DocumentExtractor(Protocol):
    """Protocol for document text extractors."""
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from a document."""
        ...


class PDFExtractor:
    """Extractor for PDF files."""
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        logger.info(f"Extracting text from PDF: {file_path.name}")
        reader = PdfReader(str(file_path))
        
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
                logger.debug(f"Extracted {len(page_text)} chars from page {page_num}")
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Successfully extracted {len(full_text)} characters from {len(reader.pages)} pages")
        
        return full_text


class DOCXExtractor:
    """Extractor for DOCX files."""
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        logger.info(f"Extracting text from DOCX: {file_path.name}")
        doc = Document(str(file_path))
        
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n\n".join(text_parts)
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        
        return full_text


class DOCExtractor:
    """Extractor for legacy DOC files (limited support)."""
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from DOC file using python-docx (limited support for old format).
        
        Args:
            file_path: Path to DOC file
            
        Returns:
            Extracted text
            
        Note:
            python-docx primarily supports .docx format. For better .doc support,
            consider using external tools or converting to .docx first.
        """
        logger.warning(f"Attempting to extract from legacy DOC format: {file_path.name}")
        logger.warning("Legacy .doc format has limited support. Consider converting to .docx")
        
        try:
            # Try using python-docx (may work for some .doc files)
            doc = Document(str(file_path))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOC")
            return full_text
        
        except Exception as e:
            logger.error(f"Failed to extract from DOC file: {str(e)}")
            raise Exception(
                f"Legacy .doc format extraction failed. "
                f"Please convert to .docx format or use a different file. Error: {str(e)}"
            )


class DocumentProcessor:
    """
    Handles document processing with support for multiple formats.
    Extensible architecture for adding new document types.
    """
    
    # Registry of supported file extensions and their extractors
    EXTRACTORS = {
        ".pdf": PDFExtractor,
        ".docx": DOCXExtractor,
        ".doc": DOCExtractor,
    }
    
    def __init__(self, chunk_size: int = 400, chunk_overlap: int = 50):
        """
        Initialize document processor.
        
        Args:
            chunk_size: Target size for text chunks (in characters)
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"DocumentProcessor initialized with chunk_size={chunk_size}, overlap={chunk_overlap}")
    
    @classmethod
    def register_extractor(cls, extension: str, extractor_class: type):
        """
        Register a new document extractor for a file extension.
        
        Args:
            extension: File extension (e.g., ".txt", ".md")
            extractor_class: Extractor class implementing DocumentExtractor protocol
        """
        cls.EXTRACTORS[extension.lower()] = extractor_class
        logger.info(f"Registered extractor for {extension}: {extractor_class.__name__}")
    
    @classmethod
    def get_supported_extensions(cls) -> List[str]:
        """
        Get list of supported file extensions.
        
        Returns:
            List of supported extensions
        """
        return list(cls.EXTRACTORS.keys())
    
    def extract_text_from_document(self, file_path: Path) -> str:
        """
        Extract text from any supported document format.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text as a single string
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
            Exception: If extraction fails
        """
        if not file_path.exists():
            logger.error(f"Document file not found: {file_path}")
            raise FileNotFoundError(f"Document file not found: {file_path}")
        
        # Get file extension
        extension = file_path.suffix.lower()
        
        # Check if format is supported
        if extension not in self.EXTRACTORS:
            logger.error(f"Unsupported file format: {extension}")
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported formats: {', '.join(self.EXTRACTORS.keys())}"
            )
        
        try:
            # Get appropriate extractor
            extractor_class = self.EXTRACTORS[extension]
            extractor = extractor_class()
            
            # Extract text
            text = extractor.extract_text(file_path)
            
            # Clean up the text
            text = self._clean_text(text)
            
            return text
            
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {str(e)}")
            raise Exception(f"Document extraction failed: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text by removing excessive whitespace and special characters.
        
        Args:
            text: Raw extracted text
            
        Returns:
            Cleaned text
        """
        # Replace multiple spaces with single space
        text = re.sub(r' +', ' ', text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n\n+', '\n\n', text)
        
        # Remove special characters that might cause issues
        text = text.replace('\x00', '')
        
        # Remove excessive whitespace at line boundaries
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        return text.strip()
    
    def chunk_text(self, text: str, metadata: dict = None) -> List[Tuple[str, dict]]:
        """
        Split text into overlapping chunks for embedding.
        
        Args:
            text: Full text to be chunked
            metadata: Optional metadata to attach to each chunk
            
        Returns:
            List of tuples containing (chunk_text, chunk_metadata)
        """
        if not text or not text.strip():
            logger.warning("Attempted to chunk empty text")
            return []
        
        metadata = metadata or {}
        chunks = []
        
        # Split by sentences to avoid breaking mid-sentence
        sentences = self._split_into_sentences(text)
        
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_length = len(sentence)
            
            # If adding this sentence exceeds chunk size and we have content, save chunk
            if current_length + sentence_length > self.chunk_size and current_chunk:
                chunk_text = " ".join(current_chunk)
                chunk_metadata = {
                    **metadata,
                    "chunk_length": len(chunk_text),
                    "chunk_index": len(chunks)
                }
                chunks.append((chunk_text, chunk_metadata))
                
                # Keep last few sentences for overlap
                overlap_text = " ".join(current_chunk)
                overlap_length = len(overlap_text)
                
                # Remove sentences from start until we're within overlap size
                while overlap_length > self.chunk_overlap and current_chunk:
                    removed = current_chunk.pop(0)
                    overlap_length -= len(removed) + 1  # +1 for space
                
                current_length = overlap_length
            
            current_chunk.append(sentence)
            current_length += sentence_length + 1  # +1 for space between sentences
        
        # Add the last chunk if there's remaining content
        if current_chunk:
            chunk_text = " ".join(current_chunk)
            chunk_metadata = {
                **metadata,
                "chunk_length": len(chunk_text),
                "chunk_index": len(chunks)
            }
            chunks.append((chunk_text, chunk_metadata))
        
        logger.info(f"Created {len(chunks)} chunks from {len(text)} characters")
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences using basic sentence boundary detection.
        
        Args:
            text: Text to split
            
        Returns:
            List of sentences
        """
        # Simple sentence splitter - splits on .!? followed by space or newline
        sentence_pattern = r'(?<=[.!?])\s+(?=[A-Z])'
        sentences = re.split(sentence_pattern, text)
        
        # Filter out empty sentences and strip whitespace
        sentences = [s.strip() for s in sentences if s.strip()]
        
        return sentences
    
    def process_document(self, file_path: Path, document_id: str) -> List[Tuple[str, dict]]:
        """
        Complete document processing pipeline: extract text and create chunks.
        
        Args:
            file_path: Path to document file
            document_id: Unique identifier for the document
            
        Returns:
            List of (chunk_text, metadata) tuples ready for embedding
        """
        logger.info(f"Processing document: {file_path.name} (ID: {document_id})")
        
        # Extract text
        text = self.extract_text_from_document(file_path)
        
        # Create metadata for all chunks
        base_metadata = {
            "document_id": document_id,
            "filename": file_path.name,
            "file_type": file_path.suffix.lower(),
            "total_length": len(text)
        }
        
        # Chunk the text
        chunks = self.chunk_text(text, metadata=base_metadata)
        
        logger.info(f"Document processing complete: {len(chunks)} chunks created")
        return chunks


# For backward compatibility - keep PDFProcessor as alias
PDFProcessor = DocumentProcessor
