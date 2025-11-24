"""
Metadata extraction module for documents.
Extracts file system metadata, document properties, and XMP metadata.
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional
from datetime import datetime
from pypdf import PdfReader
from docx import Document as DocxDocument
import os

logger = logging.getLogger(__name__)


class MetadataExtractor:
    """
    Extracts comprehensive metadata from documents.
    Supports file system metadata, document properties, and format-specific metadata.
    """
    
    @staticmethod
    def extract_file_metadata(file_path: Path) -> Dict[str, Any]:
        """
        Extract basic file system metadata.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing file metadata
        """
        try:
            stat = file_path.stat()
            
            metadata = {
                "filename": file_path.name,
                "file_path": str(file_path.absolute()),
                "file_size": stat.st_size,  # bytes
                "file_size_mb": round(stat.st_size / (1024 * 1024), 2),
                "file_type": file_path.suffix.lower(),
                "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "accessed_at": datetime.fromtimestamp(stat.st_atime).isoformat(),
            }
            
            logger.debug(f"Extracted file metadata for: {file_path.name}")
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract file metadata: {str(e)}")
            return {}
    
    @staticmethod
    def extract_pdf_metadata(file_path: Path) -> Dict[str, Any]:
        """
        Extract PDF-specific metadata including XMP data.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dictionary containing PDF metadata
        """
        try:
            reader = PdfReader(str(file_path))
            
            metadata = {
                "page_count": len(reader.pages),
                "is_encrypted": reader.is_encrypted,
            }
            
            # Extract document information
            if reader.metadata:
                info = reader.metadata
                
                # Standard PDF metadata fields
                metadata.update({
                    "title": info.get("/Title", ""),
                    "author": info.get("/Author", ""),
                    "subject": info.get("/Subject", ""),
                    "creator": info.get("/Creator", ""),
                    "producer": info.get("/Producer", ""),
                    "keywords": info.get("/Keywords", ""),
                    "creation_date": str(info.get("/CreationDate", "")),
                    "modification_date": str(info.get("/ModDate", "")),
                })
            
            # Extract XMP metadata if available
            if hasattr(reader, 'xmp_metadata') and reader.xmp_metadata:
                try:
                    xmp = reader.xmp_metadata
                    
                    xmp_data = {}
                    
                    # Try to extract common XMP fields
                    if hasattr(xmp, 'dc_title'):
                        xmp_data['xmp_title'] = xmp.dc_title
                    if hasattr(xmp, 'dc_creator'):
                        xmp_data['xmp_creator'] = xmp.dc_creator
                    if hasattr(xmp, 'dc_description'):
                        xmp_data['xmp_description'] = xmp.dc_description
                    if hasattr(xmp, 'dc_subject'):
                        xmp_data['xmp_subject'] = xmp.dc_subject
                    if hasattr(xmp, 'pdf_keywords'):
                        xmp_data['xmp_keywords'] = xmp.pdf_keywords
                    if hasattr(xmp, 'pdf_producer'):
                        xmp_data['xmp_producer'] = xmp.pdf_producer
                    
                    metadata['xmp_metadata'] = xmp_data
                    
                except Exception as e:
                    logger.debug(f"Could not extract XMP metadata: {str(e)}")
            
            logger.info(f"Extracted PDF metadata: {metadata.get('page_count', 0)} pages")
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract PDF metadata: {str(e)}")
            return {"page_count": 0}
    
    @staticmethod
    def extract_docx_metadata(file_path: Path) -> Dict[str, Any]:
        """
        Extract DOCX-specific metadata from core properties.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary containing DOCX metadata
        """
        try:
            doc = DocxDocument(str(file_path))
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            
            # Extract core properties
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                
                metadata.update({
                    "title": props.title or "",
                    "author": props.author or "",
                    "subject": props.subject or "",
                    "keywords": props.keywords or "",
                    "comments": props.comments or "",
                    "category": props.category or "",
                    "created": props.created.isoformat() if props.created else "",
                    "modified": props.modified.isoformat() if props.modified else "",
                    "last_modified_by": props.last_modified_by or "",
                    "revision": props.revision or 0,
                })
            
            logger.info(f"Extracted DOCX metadata: {metadata.get('paragraph_count', 0)} paragraphs")
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract DOCX metadata: {str(e)}")
            return {"paragraph_count": 0}
    
    @staticmethod
    def extract_doc_metadata(file_path: Path) -> Dict[str, Any]:
        """
        Extract DOC-specific metadata (limited support).
        
        Args:
            file_path: Path to DOC file
            
        Returns:
            Dictionary containing DOC metadata
        """
        try:
            # Limited support for legacy .doc format
            # Try to extract basic info using python-docx
            doc = DocxDocument(str(file_path))
            
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "legacy_format": True,
            }
            
            logger.warning(f"Limited metadata extraction for legacy .doc format")
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract DOC metadata: {str(e)}")
            return {"legacy_format": True}
    
    def extract_all_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract all available metadata for a document.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Comprehensive metadata dictionary
        """
        logger.info(f"Extracting metadata for: {file_path.name}")
        
        # Start with file system metadata
        metadata = self.extract_file_metadata(file_path)
        
        # Add format-specific metadata
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.pdf':
            format_metadata = self.extract_pdf_metadata(file_path)
        elif file_extension == '.docx':
            format_metadata = self.extract_docx_metadata(file_path)
        elif file_extension == '.doc':
            format_metadata = self.extract_doc_metadata(file_path)
        else:
            format_metadata = {}
        
        # Merge metadata
        metadata.update(format_metadata)
        
        # Add extraction timestamp
        metadata['metadata_extracted_at'] = datetime.now().isoformat()
        
        logger.info(f"Metadata extraction complete for: {file_path.name}")
        return metadata
    
    @staticmethod
    def clean_metadata_for_storage(metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Clean and normalize metadata for ChromaDB storage.
        Converts complex types to strings and handles None values.
        
        Args:
            metadata: Raw metadata dictionary
            
        Returns:
            Cleaned metadata suitable for ChromaDB
        """
        cleaned = {}
        
        for key, value in metadata.items():
            if value is None or value == "":
                continue
            
            # Convert to string or number (ChromaDB requirements)
            if isinstance(value, (str, int, float, bool)):
                cleaned[key] = value
            elif isinstance(value, dict):
                # Flatten nested dicts or convert to JSON string
                for nested_key, nested_value in value.items():
                    if isinstance(nested_value, (str, int, float, bool)):
                        cleaned[f"{key}_{nested_key}"] = nested_value
            elif isinstance(value, list):
                # Convert lists to comma-separated strings
                cleaned[key] = ", ".join(str(item) for item in value)
            else:
                # Convert everything else to string
                cleaned[key] = str(value)
        
        return cleaned
